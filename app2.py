import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import argparse

# 👉 Eger Tesseract PATH gerek bolsa aç:
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\user\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"

# ---------------- CLEAN TEXT ----------------
import re

def clean_text(text):
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


# ---------------- MAIN CLASS ----------------
class PDFProcessor:
    def __init__(self, pdf_path, lang="eng+rus"):
        self.pdf_path = pdf_path
        self.lang = lang
        self.doc = fitz.open(pdf_path)

    # ---------- TEXT + HEADING ----------
    def extract_text_with_headings(self, page):
        blocks = page.get_text("dict")["blocks"]
        result = []

        for b in blocks:
            if "lines" in b:
                for line in b["lines"]:
                    line_text = ""
                    max_size = 0

                    for span in line["spans"]:
                        line_text += span["text"] + " "
                        max_size = max(max_size, span["size"])

                    line_text = line_text.strip()

                    if not line_text:
                        continue

                    # Heading detection
                    if max_size >= 14:
                        result.append(f"## {line_text}")
                    else:
                        result.append(line_text)

        return "\n".join(result)

    # ---------- TABLE EXTRACTION ----------
    def extract_tables(self, page):
        table_text = ""

        try:
            tables = page.find_tables()

            if tables and tables.tables:
                for table in tables.tables:
                    data = table.extract()

                    table_text += "\n--- TABLE ---\n"

                    for row in data:
                        table_text += "\t".join([str(c) if c else "" for c in row]) + "\n"

                    table_text += "\n"
        except:
            pass

        return table_text

    # ---------- OCR ----------
    def ocr_page(self, page):
        pix = page.get_pixmap(dpi=300)
        img = Image.open(io.BytesIO(pix.tobytes("png")))

        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(img, lang=self.lang, config=custom_config)

        return text

    # ---------- SMART CHECK ----------
    def is_weak_text(self, text):
        return len(text.strip()) < 50

    # ---------- PROCESS ----------
    def process(self):
        final_text = []

        for i in range(len(self.doc)):
            page = self.doc[i]

            text = self.extract_text_with_headings(page)
            table = self.extract_tables(page)

            if self.is_weak_text(text):
                print(f"Page {i+1}: OCR running...")
                text = self.ocr_page(page)
            else:
                print(f"Page {i+1}: Text extracted")

            page_text = f"\n--- Page {i+1} ---\n{text}\n{table}"
            final_text.append(clean_text(page_text))

        return "\n\n".join(final_text)


# ---------------- SAVE TXT ----------------
def save_txt(text, path):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"TXT saved: {path}")


# ---------------- SAVE DOCX ----------------
def save_docx(text, path):
    from docx import Document

    doc = Document()

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        if line.startswith("--- Page"):
            doc.add_heading(line, level=1)

        elif line.startswith("##"):
            doc.add_heading(line.replace("##", ""), level=2)

        elif line.startswith("--- TABLE ---"):
            doc.add_paragraph("Table:")

        elif "\t" in line:
            doc.add_paragraph(" | ".join(line.split("\t")))

        else:
            doc.add_paragraph(line)

    doc.save(path)
    print(f"DOCX saved: {path}")


# ---------------- CLI ----------------
def run_cli():
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("--txt")
    parser.add_argument("--docx")
    parser.add_argument("--lang", default="eng+rus")

    args = parser.parse_args()

    processor = PDFProcessor(args.input, args.lang)
    text = processor.process()

    if args.txt:
        save_txt(text, args.txt)

    if args.docx:
        save_docx(text, args.docx)

    if not args.txt and not args.docx:
        print(text)


# ---------------- GUI ----------------
def run_gui():
    import tkinter as tk
    from tkinter import filedialog, messagebox

    def select_file():
        path = filedialog.askopenfilename(filetypes=[("PDF", "*.pdf")])
        entry.delete(0, tk.END)
        entry.insert(0, path)

    def convert():
        path = entry.get()

        if not path:
            messagebox.showerror("Error", "Select PDF")
            return

        processor = PDFProcessor(path)
        text = processor.process()

        txt_path = path.replace(".pdf", ".txt")
        docx_path = path.replace(".pdf", ".docx")

        save_txt(text, txt_path)
        save_docx(text, docx_path)

        messagebox.showinfo("Done", "Converted successfully!")

    root = tk.Tk()
    root.title("PDF Pro OCR Tool")
    root.geometry("400x180")

    entry = tk.Entry(root, width=40)
    entry.pack(pady=10)

    tk.Button(root, text="Select PDF", command=select_file).pack()
    tk.Button(root, text="Convert", command=convert).pack(pady=10)

    root.mainloop()


# ---------------- MAIN ----------------
if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        run_cli()
    else:
        run_gui()